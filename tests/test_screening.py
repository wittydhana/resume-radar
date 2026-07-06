import io
import tempfile
import unittest
from contextlib import redirect_stdout
from pathlib import Path

from docx import Document
from fpdf import FPDF

from pipeline import run_screening


class ScreeningPipelineTests(unittest.TestCase):
    def test_run_screening_writes_outputs_and_prints_summary(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            base = Path(tmpdir)
            jd_path = base / "job_description.txt"
            resumes_dir = base / "resumes"
            out_dir = base / "output"
            resumes_dir.mkdir(parents=True, exist_ok=True)

            jd_text = "Python, SQL, machine learning, NLP, pandas, scikit-learn"
            jd_path.write_text(jd_text, encoding="utf-8")

            (resumes_dir / "candidate_a.txt").write_text(
                "Python developer with SQL and pandas experience in NLP projects",
                encoding="utf-8"
            )

            docx_path = resumes_dir / "candidate_b.docx"
            doc = Document()
            doc.add_paragraph("Candidate B")
            doc.add_paragraph("Experienced in Python, machine learning, and scikit-learn.")
            doc.save(docx_path)

            pdf_path = resumes_dir / "candidate_c.pdf"
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, "Candidate C", ln=True)
            pdf.cell(0, 10, "Worked on SQL and NLP projects.", ln=True)
            pdf.output(str(pdf_path))

            stdout = io.StringIO()
            with redirect_stdout(stdout):
                df = run_screening(jd_path, resumes_dir, out_dir)

            self.assertEqual(len(df), 3)
            self.assertTrue((out_dir / "ranked_results.csv").exists())
            self.assertTrue((out_dir / "ranked_results.json").exists())
            self.assertIn("Ranked", stdout.getvalue())
            self.assertTrue((out_dir / "scoring_note.txt").exists())
            self.assertIn("excerpt", df.columns)


if __name__ == "__main__":
    unittest.main()
