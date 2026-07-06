from pathlib import Path
from docx import Document
from fpdf import FPDF

SAMPLES = Path("samples/resumes")
SAMPLES.mkdir(parents=True, exist_ok=True)

profiles = [
    {
        "filename": "vinay_patel.docx",
        "name": "Vinay Patel",
        "title": "AI Research Intern",
        "summary": "Experienced in Python, NLP, and machine learning with hands-on project work.",
        "experience": [
            "Built text classification models using Python, scikit-learn, and pandas.",
            "Extracted features from unstructured data and tuned model performance.",
            "Implemented NLP pipelines for sentiment analysis and entity extraction."
        ],
        "skills": "Python, SQL, NLP, machine learning, pandas, scikit-learn, AWS"
    },
    {
        "filename": "nisha_varma.pdf",
        "name": "Nisha Varma",
        "title": "Machine Learning Engineer",
        "summary": "Developer with experience using Python and SQL to build production ML systems.",
        "experience": [
            "Designed data pipelines and feature stores using Python and SQL.",
            "Developed classification models and evaluated using sklearn metrics.",
            "Collaborated with product teams to deploy small AI prototypes."
        ],
        "skills": "Python, SQL, machine learning, scikit-learn, pandas, Docker"
    },
    {
        "filename": "ananya_mishra.txt",
        "name": "Ananya Mishra",
        "title": "Data Science Intern",
        "summary": "Data science intern with coursework in machine learning, NLP, and analytics.",
        "experience": [
            "Built exploratory data analysis reports and dashboards.",
            "Applied classification models using sklearn and pandas.",
            "Worked with SQL databases to clean and join datasets."
        ],
        "skills": "Python, SQL, pandas, machine learning, NLP, statistics"
    },
    {
        "filename": "ravi_gupta.docx",
        "name": "Ravi Gupta",
        "title": "Business Intelligence Analyst",
        "summary": "BI analyst focused on SQL, reporting, and translating business needs into data insights.",
        "experience": [
            "Created dashboards using SQL and Excel for marketing reporting.",
            "Presented data-driven product recommendations to leadership.",
            "Supported analytics for revenue and customer retention."
        ],
        "skills": "SQL, Excel, data visualization, business analysis, communication"
    },
    {
        "filename": "mina_bose.pdf",
        "name": "Mina Bose",
        "title": "NLP Engineer",
        "summary": "NLP engineer with experience building text pipelines and working with transformer models.",
        "experience": [
            "Built document classifiers and named entity recognition systems.",
            "Evaluated NLP models using precision, recall, and F1 score.",
            "Collaborated with engineering teams to deploy model APIs."
        ],
        "skills": "Python, NLP, machine learning, transformers, API"
    },
    {
        "filename": "kavita_mehta.txt",
        "name": "Kavita Mehta",
        "title": "Analytics Engineer",
        "summary": "Analytics engineer with strong SQL, Python scripting, and data pipeline experience.",
        "experience": [
            "Built ETL processes using Python and SQL for customer analytics.",
            "Designed reusable data transformations and validation checks.",
            "Worked with stakeholders to define metrics and reporting needs."
        ],
        "skills": "Python, SQL, data engineering, analytics, pandas"
    },
    {
        "filename": "sahil_sharma.docx",
        "name": "Sahil Sharma",
        "title": "Junior Data Scientist",
        "summary": "Junior data scientist experienced in model development and prototyping AI solutions.",
        "experience": [
            "Built regression and classification models using scikit-learn.",
            "Performed feature engineering on structured and text data.",
            "Created model evaluation dashboards for stakeholders."
        ],
        "skills": "Python, scikit-learn, machine learning, pandas, data visualization"
    },
    {
        "filename": "tara_narayanan.pdf",
        "name": "Tara Narayanan",
        "title": "Software Engineer",
        "summary": "Software engineer with experience building backend APIs and integrating data workflows.",
        "experience": [
            "Built API services using Python and FastAPI.",
            "Integrated data ingestion jobs with SQL and cloud storage.",
            "Wrote unit tests and documentation for data tools."
        ],
        "skills": "Python, API, FastAPI, SQL, Docker"
    },
    {
        "filename": "arjun_rao.txt",
        "name": "Arjun Rao",
        "title": "Graduate Research Assistant",
        "summary": "Graduate researcher working on machine learning experiments and academic NLP projects.",
        "experience": [
            "Conducted NLP research using Python and scientific computing.",
            "Published analyses on text classification and model interpretability.",
            "Developed reproducible notebooks and shared results with advisors."
        ],
        "skills": "Python, NLP, machine learning, research, pandas"
    },
    {
        "filename": "leah_patel.docx",
        "name": "Leah Patel",
        "title": "Product Analyst",
        "summary": "Product analyst who combines data insights with product strategy and stakeholder communication.",
        "experience": [
            "Analyzed feature adoption using SQL and product metrics.",
            "Worked with engineering teams to prioritize roadmap items.",
            "Presented findings to cross-functional teams."
        ],
        "skills": "SQL, analytics, communication, product strategy, Python"
    }
]

for profile in profiles:
    path = SAMPLES / profile["filename"]
    if path.suffix.lower() == ".docx":
        doc = Document()
        doc.add_heading(profile["name"], level=1)
        doc.add_paragraph(profile["title"])
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(profile["summary"])
        doc.add_heading("Experience", level=2)
        for item in profile["experience"]:
            doc.add_paragraph(item, style="List Bullet")
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(profile["skills"])
        doc.save(path)
    elif path.suffix.lower() == ".pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, profile["name"], ln=True)
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 8, profile["title"], ln=True)
        pdf.ln(4)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, "Summary", ln=True)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 8, profile["summary"])
        pdf.ln(2)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, "Experience", ln=True)
        pdf.set_font("Arial", "", 12)
        for item in profile["experience"]:
            pdf.multi_cell(0, 8, f"- {item}")
        pdf.ln(2)
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, "Skills", ln=True)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 8, profile["skills"])
        pdf.output(str(path))
    elif path.suffix.lower() == ".txt":
        with path.open("w", encoding="utf-8") as fh:
            fh.write(profile["name"] + "\n")
            fh.write(profile["title"] + "\n\n")
            fh.write("Summary:\n")
            fh.write(profile["summary"] + "\n\n")
            fh.write("Experience:\n")
            for item in profile["experience"]:
                fh.write(f"- {item}\n")
            fh.write("\nSkills:\n")
            fh.write(profile["skills"] + "\n")

print(f"Generated sample resumes in {SAMPLES}")
